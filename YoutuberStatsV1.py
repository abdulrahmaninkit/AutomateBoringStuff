import bs4
import requests
import re
import docx
import os
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Get the output directory from the environment variables
output_dir = os.getenv('DOWNLOAD_PATH')

def getinfo(name):
    #Connecting to website.
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    res = requests.get(name, headers=headers)
    res.raise_for_status()
    #Code of Extration of information.
    soup = bs4.BeautifulSoup(res.text, 'html.parser')
    

    #For creating the PDF file.
    pdf_path = os.path.join(output_dir, f"{cname}.pdf")
    document = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables = []

    #Adding the heading

    heading_text = f"{cname} Report"
    heading = Paragraph(heading_text, styles["Heading1"])
    flowables.append(heading)

    #For creating the word file
    word_path = os.path.join(output_dir, f"{cname}.docx")
    doc = docx.Document()
    doc.add_heading('Youtubers Data')
    name = (f"Youtubers Name: {cname}")
    doc.add_heading(name)

    #paths of the element that is being extracted.
    elems1 = soup.select(f'#YouTubeUserTopInfoBlock > div:nth-child({2})')
    elems2 = soup.select(f'#YouTubeUserTopInfoBlock > div:nth-child({3})')
    elems3 = soup.select(f'#YouTubeUserTopInfoBlock > div:nth-child({4})')
    elems4 = soup.select(f'#YouTubeUserTopInfoBlock > div:nth-child({5})')
    elems5 = soup.select(f'#YouTubeUserTopInfoBlock > div:nth-child({6})')
    elems6 = soup.select(f'#YouTubeUserTopInfoBlock > div:nth-child(7) > span:nth-child(3)')

    components1 = elems1[0].text.split()
    components2 = elems2[0].text.split()
    components3 = elems3[0].text.split()
    components4 = elems4[0].text.split()
    components5 = elems5[0].text.split()
    components6 = elems6[0].text.split()

    date = components6[0]+' '+components6[1]+' '+components6[2]
    text1 = (f""" 
                Total Uploads: {components1[1]}
                Subscribers: {components2[1]}
                Video Views: {components3[2]}
                Contry of Channel: {components4[1]}
                Channel Type: {components5[2]}
                Date of Creation: {date}
            """)
    
    #Adding to the word file
    doc.add_paragraph(text1)

    #Adding to the PDF file
    paragraph1 = Paragraph(text1, styles["Normal"])
    flowables.append(paragraph1)
    flowables.append(Spacer(1, 12)) 
    flowables.append(Spacer(1, 12))

    #Path of the element that is being extracted.
    elems = soup.select('#socialblade-user-content > div:nth-child(1) > div:nth-child(2)')
    #Trying to extract only the usefull information.
    if elems:
        numbers = elems[0].text.strip()
        numRegex = re.compile(r'''
            # 53,000 or 199 or 1,947 or 55 or 2
                (
                   (\d{1,3})?
                   (,)?
                   (\d{3}|\d{2}|\d{1})           
                )           
        ''', re.VERBOSE)
        extractedNum = numRegex.findall(numbers)
        allNumbers = []

        for Numbers in extractedNum:
            allNumbers.append(Numbers[0])
        
        text2 = (f""" 
                Social Rank: {allNumbers[0]}
                Subscribe Rank: {allNumbers[1]}
                Country Rank: {allNumbers[2]}
                Domain Rank: {allNumbers[3]}
                """)
        #Adding to the word file
        doc.add_paragraph(text2)

        #Adding to the PDF file.
        paragraph2 = Paragraph(text2, styles["Normal"])
        flowables.append(paragraph2)
        flowables.append(Spacer(1, 12))
        flowables.append(Spacer(1, 12))
    data = []
    #Getting the stats info
    for i in range(7,20):
        elems = soup.select(f'#socialblade-user-content > div:nth-child({i})')
        components = elems[0].text.split()
        
        entry={
            "Date" : components[0],
            "Day" : components[1],
            "Subscribers_change" : components[2],
            "Total_subscribers" : components[3],
            "Views_change" : components[4],
            "Total_views" : components[5],
            "Earningsl" : components[6],
            "Earningsu" : components[8]
        }
        data.append(entry)
        

        #Saving the data into a proper formatat that can be displayed
        for entry in data:
            text3 = (f"""Date: {entry['Date']}
                        Day: {entry['Day']} 

                        Subscribers Change: {entry['Subscribers_change']} 
                        Total Subscribers: {entry['Total_subscribers']} 

                        Views Change: {entry['Views_change']}
                        Total Views: {entry['Total_views']} 

                        Estimated Earnings: {entry['Earningsl']} - {entry['Earningsu']}""")
            
        #Adding to the PDF file.
        paragraph3 = Paragraph(text3, styles["Normal"])
        flowables.append(paragraph3)
        flowables.append(Spacer(1, 12))

        #Adding the content to the word file.
        doc.add_paragraph(text3)

    document.build(flowables)

    #Saving the word file into the local.

    doc.save(word_path) 
    
    #Saving the PDF file into the local.
    #pdf.output(f"{cname}.pdf")

    #Saving the data into the excel file.

    #Creating a workbook
    wb = Workbook()

    #Selecting the active sheet
    sheet = wb.active

    #Saving the data

    sheet["D1"] = f"{cname}"
    sheet["E1"] = "CHANNEL"
    sheet["F1"] = "DATA"

    sheet['A2'] = "Date"
    sheet['B2'] = "Day"
    sheet['C2'] = "Subscribers Change"
    sheet['D2'] = "Total Subscribers"
    sheet['E2'] = "Views Change"
    sheet['F2'] = "Total Views"
    sheet['G2'] = "Estimated Earnings"

    # for index, entry in enumerate(data, start=2):
    #     sheet[f'A{index}'] = entry['Date']
    #     sheet[f'B{index}'] = entry['Day']
    #     sheet[f'C{index}'] = entry['Subscribers_change']
    #     sheet[f'D{index}'] = entry['Total_subscribers']
    #     sheet[f'E{index}'] = entry['Views_change']
    #     sheet[f'F{index}'] = entry['Total_views']
    #     sheet[f'G{index}'] = f"{entry['Earningsl']} - {entry['Earningsu']}"

    index=3
    for entry in data:
        sheet[f'A{index}'] = entry['Date']
        sheet[f'B{index}'] = entry['Day']
        sheet[f'C{index}'] = entry['Subscribers_change']
        sheet[f'D{index}'] = entry['Total_subscribers']
        sheet[f'E{index}'] = entry['Views_change']
        sheet[f'F{index}'] = entry['Total_views']
        sheet[f'G{index}'] = f"{entry['Earningsl']} - {entry['Earningsu']}"
        index = index+1
    # Save the workbook
    excel_path = os.path.join(output_dir, f"{cname}.xlsx")
    wb.save(excel_path)
    
cname=input("")
getinfo(f'https://socialblade.com/youtube/c/{cname}_')



