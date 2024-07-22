import docx
d = docx.Document("Ticketing System Design.docx")

print(d.paragraphs[0].text)
print(d.paragraphs[3].text)

p = d.paragraphs[3]
print(p.runs[0].text)

p.runs[0].text = 'italic and underline'
d.save("Ticketing System Design.docx")

d.add_paragraph('''

The file is being modified using docx module.
This module can be used to read,write doc files 

''')
d.save("Ticketing System Design.docx")



